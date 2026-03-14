"""
report_generator.py
-------------------
Generates a professional DOCX DDR report with:
  - Cover page
  - Executive summary
  - Thermal data table + images
  - Area-wise findings with supporting photos
  - Checklist results
  - Recommendations
  - Conclusion
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ─────────────────────────────────────────────
# Colour Palette
# ─────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x0C, 0x38, 0x70)   # #0C3870
BLUE       = RGBColor(0x1A, 0x56, 0xA0)   # #1A56A0
BLUE_LT    = RGBColor(0xE8, 0xF1, 0xFB)   # #E8F1FB
AMBER      = RGBColor(0xE0, 0x7B, 0x00)   # #E07B00
AMBER_LT   = RGBColor(0xFF, 0xF3, 0xE0)
RED        = RGBColor(0xC0, 0x39, 0x2B)   # #C0392B
RED_LT     = RGBColor(0xFD, 0xEC, 0xEA)
GREEN      = RGBColor(0x1A, 0x7A, 0x45)   # #1A7A45
GREEN_LT   = RGBColor(0xE8, 0xF5, 0xEE)
GRAY       = RGBColor(0x4A, 0x4A, 0x4A)
GRAY_LT    = RGBColor(0xF4, 0xF4, 0xF4)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x1A, 0x1A, 0x1A)


# ─────────────────────────────────────────────
# Utilities
# ─────────────────────────────────────────────

def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_border(cell, sides=("top","bottom","left","right"), size=6, color="D5D5D5"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders_el = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:color"), color)
        borders_el.append(border)
    tcPr.append(borders_el)


def _add_paragraph(doc, text="", bold=False, size=11, color=None, alignment=None,
                   space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p


def _add_heading(doc, text, level=1, color=BLUE_DARK, size=None):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size or sizes.get(level, 12))
    run.font.color.rgb = color
    # Add bottom border for H1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "1A56A0")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def _severity_color(severity: str):
    s = severity.lower()
    if s in ("high", "critical", "immediate"):
        return RED, RED_LT
    elif s in ("medium", "moderate", "within 30 days"):
        return AMBER, AMBER_LT
    elif s in ("low", "within 90 days", "monitor"):
        return GREEN, GREEN_LT
    else:
        return BLUE, BLUE_LT


def _insert_images_row(doc, image_paths: list, max_per_row=3, width_inches=1.9, caption=""):
    """Insert a row of images with optional caption."""
    valid = [p for p in image_paths if p and p != "Image Not Available" and os.path.exists(p)]
    if not valid:
        _add_paragraph(doc, "  [Image Not Available]", italic=True, color=GRAY, size=10)
        return

    # Process in chunks of max_per_row
    for chunk_start in range(0, len(valid), max_per_row):
        chunk = valid[chunk_start:chunk_start + max_per_row]
        table = doc.add_table(rows=1, cols=len(chunk))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"

        for col_idx, img_path in enumerate(chunk):
            cell = table.cell(0, col_idx)
            _set_cell_border(cell)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = cell.paragraphs[0].add_run()
                run.add_picture(img_path, width=Inches(width_inches))
            except Exception:
                cell.paragraphs[0].add_run("[Image Error]")

        doc.add_paragraph()  # spacing after table

    if caption:
        _add_paragraph(doc, caption, italic=True, color=GRAY, size=9, space_before=0, space_after=8)


# ─────────────────────────────────────────────
# Main Generator
# ─────────────────────────────────────────────

def generate_report(inspection_data, thermal_data, analysis: dict, output_path: str) -> str:
    """
    Generates the full DDR DOCX report.
    Returns the output file path.
    """
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.page_width   = Cm(21)
        section.page_height  = Cm(29.7)
        section.left_margin  = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin   = Cm(2)
        section.bottom_margin = Cm(2)

    _build_cover(doc, inspection_data, analysis)
    _build_executive_summary(doc, inspection_data, thermal_data, analysis)
    _build_thermal_section(doc, thermal_data)
    _build_area_findings(doc, inspection_data, thermal_data, analysis)
    _build_checklist_section(doc, inspection_data)
    _build_recommendations(doc, analysis)
    _build_conclusion(doc, analysis)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


# ─────────────────────────────────────────────
# Section Builders
# ─────────────────────────────────────────────

def _build_cover(doc, inspection_data, analysis):
    """Cover page."""
    doc.add_paragraph()
    doc.add_paragraph()

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DETAILED DEFECT REPORT")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BLUE_DARK

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Water Ingress & Dampness Investigation")
    run2.font.size = Pt(14)
    run2.font.color.rgb = GRAY

    doc.add_paragraph()

    # Divider
    p3 = _add_paragraph(doc, "─" * 60, color=BLUE, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # Property info table
    info_table = doc.add_table(rows=0, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def add_info_row(label, value):
        row = info_table.add_row()
        c1, c2 = row.cells
        _set_cell_bg(c1, GRAY_LT)
        _set_cell_border(c1)
        _set_cell_border(c2)
        c1.paragraphs[0].add_run(label).bold = True
        c1.paragraphs[0].runs[0].font.size = Pt(11)
        c2.paragraphs[0].add_run(value).font.size = Pt(11)
        c1.width = Inches(2.5)
        c2.width = Inches(3.5)

    add_info_row("Inspection Date",   inspection_data.inspection_date)
    add_info_row("Inspected By",      inspection_data.inspected_by)
    add_info_row("Property Type",     inspection_data.property_type)
    add_info_row("No. of Floors",     inspection_data.floors)
    add_info_row("Property Age",      inspection_data.property_age)
    add_info_row("Previous Audit",    inspection_data.previous_audit)
    add_info_row("Previous Repairs",  inspection_data.previous_repair)
    add_info_row("Inspection Score",  inspection_data.overall_score)
    add_info_row("Flagged Items",     inspection_data.flagged_items)
    add_info_row("Report Generated",  datetime.now().strftime("%d %B %Y"))
    add_info_row("Prepared By",       "UrbanRoof Inspection Team")

    doc.add_paragraph()
    doc.add_paragraph()

    # Overall severity badge
    sev = analysis.get("severity_assessment", "Not Available")
    sev_color, _ = _severity_color(sev.split("—")[0].strip() if "—" in sev else "Medium")
    p_sev = _add_paragraph(doc, f"Overall Severity: {sev}", bold=True, color=sev_color,
                            size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()


def _build_executive_summary(doc, inspection_data, thermal_data, analysis):
    _add_heading(doc, "1. Executive Summary", level=1)

    summary = analysis.get("executive_summary", "Not Available")
    _add_paragraph(doc, summary, size=11, space_before=4, space_after=8)

    # Metrics table
    _add_heading(doc, "Key Metrics", level=3, color=BLUE)

    metrics = [
        ("Total Areas Affected",  str(len(inspection_data.impacted_areas))),
        ("Thermal Readings",      str(len(thermal_data.readings))),
        ("Max Hotspot (°C)",      str(thermal_data.max_hotspot)),
        ("Min Coldspot (°C)",     str(thermal_data.min_coldspot)),
        ("Avg Temp Differential", f"{thermal_data.avg_delta}°C"),
        ("Thermal Anomalies",     str(thermal_data.anomaly_count)),
        ("Inspection Score",      inspection_data.overall_score),
        ("Flagged Items",         inspection_data.flagged_items),
    ]

    tbl = doc.add_table(rows=2, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    idx = 0
    for row_cells in tbl.rows:
        for cell in row_cells.cells:
            if idx < len(metrics):
                label, value = metrics[idx]
                _set_cell_bg(cell, BLUE_LT)
                _set_cell_border(cell, color="1A56A0")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(value + "\n").bold = True
                p.runs[0].font.size = Pt(16)
                p.runs[0].font.color.rgb = BLUE_DARK
                r2 = p.add_run(label)
                r2.font.size = Pt(9)
                r2.font.color.rgb = GRAY
                idx += 1

    doc.add_paragraph()

    # Impacted rooms
    _add_heading(doc, "Impacted Areas / Rooms", level=3, color=BLUE)
    rooms_str = ", ".join(inspection_data.impacted_rooms) if inspection_data.impacted_rooms else "Not Available"
    _add_paragraph(doc, rooms_str, size=11)

    # Root cause
    _add_heading(doc, "Root Cause Analysis", level=3, color=BLUE)
    rca = analysis.get("root_cause_analysis", "Not Available")
    _add_paragraph(doc, rca, size=11)


def _build_thermal_section(doc, thermal_data):
    doc.add_page_break()
    _add_heading(doc, "2. Thermal Imaging Analysis", level=1)

    interp = "Thermal imaging was performed using the Bosch GTC 400 C Professional camera (Serial No. 02700034772)."
    _add_paragraph(doc, interp, size=11, space_after=8)

    # Summary stats
    p = doc.add_paragraph()
    stats = [
        ("Readings taken", len(thermal_data.readings)),
        ("Max hotspot", f"{thermal_data.max_hotspot} °C"),
        ("Min coldspot", f"{thermal_data.min_coldspot} °C"),
        ("Average Δ", f"{thermal_data.avg_delta} °C"),
        ("Anomalies (Δ ≥ 4°C)", thermal_data.anomaly_count),
    ]
    for label, val in stats:
        run = p.add_run(f"  {label}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(f"{val}    ").font.size = Pt(10)

    doc.add_paragraph()

    # Thermal readings table
    _add_heading(doc, "Thermal Readings Summary", level=3, color=BLUE)

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = ["#", "Image ID", "Date", "Hotspot (°C)", "Coldspot (°C)", "Δ Temp (°C)"]
    for i, cell in enumerate(tbl.rows[0].cells):
        _set_cell_bg(cell, BLUE_DARK)
        _set_cell_border(cell, color="1A56A0")
        run = cell.paragraphs[0].add_run(headers[i])
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for reading in thermal_data.readings:
        row = tbl.add_row()
        delta_color, delta_bg = _severity_color(
            "High" if reading.delta >= 5 else "Medium" if reading.delta >= 3 else "Low"
        )
        values = [
            str(reading.page_num),
            reading.image_id,
            reading.date,
            str(reading.hotspot),
            str(reading.coldspot),
            f"{reading.delta}",
        ]
        for j, cell in enumerate(row.cells):
            _set_cell_border(cell)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cell.paragraphs[0].add_run(values[j])
            r.font.size = Pt(9)
            if j == 5:  # delta column — color-code
                r.bold = True
                r.font.color.rgb = delta_color

    doc.add_paragraph()

    # Thermal images — show first 6 pairs (thermal + visual)
    _add_heading(doc, "Thermal Image Gallery (Selected)", level=3, color=BLUE)
    _add_paragraph(doc, "Each pair shows the thermal image (left) and the corresponding visual photograph (right).", 
                   italic=True, size=10, color=GRAY)
    doc.add_paragraph()

    shown = 0
    for reading in thermal_data.readings:
        if shown >= 10:  # limit to 10 pairs in report
            break
        t_path = reading.thermal_image_path
        v_path = reading.visual_image_path
        t_ok = t_path != "Image Not Available" and os.path.exists(t_path)
        v_ok = v_path != "Image Not Available" and os.path.exists(v_path)

        if not (t_ok or v_ok):
            continue

        img_table = doc.add_table(rows=1, cols=2)
        img_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        img_table.style = "Table Grid"

        for col_i, (path, ok, label) in enumerate([
            (t_path, t_ok, "Thermal Image"),
            (v_path, v_ok, "Visual Photo"),
        ]):
            cell = img_table.cell(0, col_i)
            _set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if ok:
                try:
                    run = p.add_run()
                    run.add_picture(path, width=Inches(2.8))
                except Exception:
                    p.add_run("[Image Error]")
            else:
                p.add_run("[Image Not Available]").font.color.rgb = GRAY

        # Caption row
        caption_row = img_table.add_row()
        c = caption_row.cells[0].merge(caption_row.cells[1])
        _set_cell_bg(c, GRAY_LT)
        cp = c.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_c = cp.add_run(f"Image: {reading.image_id}  |  Date: {reading.date}  |  "
                           f"Hotspot: {reading.hotspot}°C  |  Coldspot: {reading.coldspot}°C  |  "
                           f"Δ = {reading.delta}°C  |  Emissivity: {reading.emissivity}")
        run_c.font.size = Pt(8)
        run_c.font.color.rgb = GRAY

        doc.add_paragraph()
        shown += 1


def _build_area_findings(doc, inspection_data, thermal_data, analysis):
    doc.add_page_break()
    _add_heading(doc, "3. Area-Wise Findings", level=1)

    ai_findings = {f["id"]: f for f in analysis.get("findings", [])}

    for area in inspection_data.impacted_areas:
        fid = f"F-{area.area_number:02d}"
        af  = ai_findings.get(fid, {})

        # Area header
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(14)
        p_hdr.paragraph_format.space_after  = Pt(2)
        run = p_hdr.add_run(f"Finding {fid}: {area.negative_description}")
        run.bold = True
        run.font.size = Pt(12)

        severity  = af.get("severity", "Medium")
        sev_color, sev_bg = _severity_color(severity)

        # Severity + priority chips
        p_chips = doc.add_paragraph()
        p_chips.paragraph_format.space_after = Pt(6)
        for label, val, col in [
            ("Severity",  severity,                        sev_color),
            ("Priority",  af.get("priority", "Not Available"), BLUE),
            ("Thermal",   af.get("thermal_evidence", "N/A"), GRAY),
        ]:
            p_chips.add_run(f" {label}: ").font.color.rgb = GRAY
            r = p_chips.add_run(f"{val}  ")
            r.bold = True
            r.font.color.rgb = col
            r.font.size = Pt(10)

        # Observation / Source 2-col table
        det_table = doc.add_table(rows=0, cols=2)
        det_table.style = "Table Grid"
        det_table.alignment = WD_TABLE_ALIGNMENT.LEFT

        def add_detail_row(label, value, bg=GRAY_LT):
            row = det_table.add_row()
            c1, c2 = row.cells
            _set_cell_bg(c1, bg)
            _set_cell_border(c1)
            _set_cell_border(c2)
            r1 = c1.paragraphs[0].add_run(label)
            r1.bold = True
            r1.font.size = Pt(10)
            c2.paragraphs[0].add_run(value).font.size = Pt(10)
            c1.width = Inches(2.0)
            c2.width = Inches(4.5)

        add_detail_row("Observation (Damage Side)",  area.negative_description)
        add_detail_row("Source Identified",          area.positive_description)
        if af.get("recommendation"):
            add_detail_row("Recommended Action",     af["recommendation"])

        doc.add_paragraph()

        # ── Negative Side Photos ──
        if area.negative_photos:
            _add_paragraph(doc, "Photographs – Damage / Affected Area:", bold=True, size=10, color=RED, space_before=4, space_after=2)
            _insert_images_row(doc, area.negative_photos, max_per_row=3, width_inches=1.85,
                               caption=f"Area {area.area_number} – Damage observed")
        else:
            _add_paragraph(doc, "Photographs – Damage Side: Image Not Available", italic=True, size=10, color=GRAY)

        # ── Positive Side Photos ──
        if area.positive_photos:
            _add_paragraph(doc, "Photographs – Source / Root Cause:", bold=True, size=10, color=AMBER, space_before=4, space_after=2)
            _insert_images_row(doc, area.positive_photos, max_per_row=3, width_inches=1.85,
                               caption=f"Area {area.area_number} – Source of water ingress")
        else:
            _add_paragraph(doc, "Photographs – Source Side: Image Not Available", italic=True, size=10, color=GRAY)

        # ── Correlated Thermal Reading ──
        idx = area.area_number - 1
        if idx < len(thermal_data.readings):
            r = thermal_data.readings[idx]
            _add_paragraph(doc, "Correlated Thermal Reading:", bold=True, size=10, color=BLUE, space_before=4, space_after=2)

            t_table = doc.add_table(rows=1, cols=4)
            t_table.style = "Table Grid"
            t_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            th_vals = [
                ("Image ID",    r.image_id),
                ("Hotspot",     f"{r.hotspot} °C"),
                ("Coldspot",    f"{r.coldspot} °C"),
                ("Δ Temp",      f"{r.delta} °C"),
            ]
            for ti, (lbl, val) in enumerate(th_vals):
                cell = t_table.rows[0].cells[ti]
                _set_cell_bg(cell, BLUE_LT)
                _set_cell_border(cell, color="1A56A0")
                cp = cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.add_run(val + "\n").bold = True
                cp.runs[0].font.size = Pt(12)
                cp.runs[0].font.color.rgb = BLUE_DARK
                r2 = cp.add_run(lbl)
                r2.font.size = Pt(8)
                r2.font.color.rgb = GRAY

            doc.add_paragraph()

            # Thermal image pair for this area
            if r.thermal_image_path != "Image Not Available" and os.path.exists(r.thermal_image_path):
                img_t = doc.add_table(rows=1, cols=2)
                img_t.style = "Table Grid"
                for ci, (ipath, label) in enumerate([
                    (r.thermal_image_path, "Thermal Image"),
                    (r.visual_image_path,  "Visual Photo"),
                ]):
                    cell = img_t.rows[0].cells[ci]
                    _set_cell_border(cell)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if ipath and ipath != "Image Not Available" and os.path.exists(ipath):
                        try:
                            p.add_run().add_picture(ipath, width=Inches(2.5))
                        except Exception:
                            p.add_run("[Image Error]")
                    else:
                        p.add_run("Image Not Available").font.color.rgb = GRAY

                cap_row = img_t.add_row()
                cm = cap_row.cells[0].merge(cap_row.cells[1])
                _set_cell_bg(cm, GRAY_LT)
                cp2 = cm.paragraphs[0]
                cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cp2.add_run(f"Thermal scan – {r.image_id}  |  Hotspot: {r.hotspot}°C  |  Coldspot: {r.coldspot}°C  |  Δ = {r.delta}°C")
                cap_run.font.size = Pt(8)
                cap_run.font.color.rgb = GRAY
                doc.add_paragraph()

        doc.add_paragraph()  # spacer between areas


def _build_checklist_section(doc, inspection_data):
    doc.add_page_break()
    _add_heading(doc, "4. Inspection Checklist Results", level=1)

    _add_paragraph(doc, "The following checklist items were evaluated during the site inspection:", size=11, space_after=8)

    # Group by category
    categories = {}
    for item in inspection_data.checklist_items:
        categories.setdefault(item.category, []).append(item)

    for category, items in categories.items():
        _add_heading(doc, category, level=3, color=BLUE)

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header
        for ci, hdr in enumerate(["Checklist Item", "Finding"]):
            cell = tbl.rows[0].cells[ci]
            _set_cell_bg(cell, BLUE_DARK)
            _set_cell_border(cell, color="1A56A0")
            r = cell.paragraphs[0].add_run(hdr)
            r.bold = True
            r.font.color.rgb = WHITE
            r.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for item in items:
            row = tbl.add_row()
            c1, c2 = row.cells
            _set_cell_border(c1)
            _set_cell_border(c2)
            c1.paragraphs[0].add_run(item.item).font.size = Pt(10)
            c1.width = Inches(4.0)
            c2.width = Inches(2.5)

            # Color-code value
            val = item.value
            val_run = c2.paragraphs[0].add_run(val)
            val_run.bold = True
            val_run.font.size = Pt(10)
            val_lower = val.lower()
            if val_lower in ("yes", "all time"):
                val_run.font.color.rgb = RED
            elif val_lower in ("moderate",):
                val_run.font.color.rgb = AMBER
            elif val_lower in ("no", "n/a"):
                val_run.font.color.rgb = GRAY
            c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()


def _build_recommendations(doc, analysis):
    doc.add_page_break()
    _add_heading(doc, "5. Recommendations", level=1)

    _add_paragraph(doc, "The following remediation actions are recommended, listed in order of priority:", size=11, space_after=8)

    recs = analysis.get("recommendations", [])
    if not recs:
        _add_paragraph(doc, "Not Available", italic=True, color=GRAY, size=11)
        return

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ci, hdr in enumerate(["#", "Recommended Action", "Area", "Priority"]):
        cell = tbl.rows[0].cells[ci]
        _set_cell_bg(cell, BLUE_DARK)
        _set_cell_border(cell, color="1A56A0")
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    widths = [Inches(0.4), Inches(4.0), Inches(1.5), Inches(1.0)]

    for i, rec in enumerate(recs):
        row = tbl.add_row()
        cells = row.cells
        for ci, (width, val) in enumerate(zip(widths, [
            str(i+1),
            rec.get("action", "Not Available"),
            rec.get("area",   "Not Available"),
            rec.get("priority", "Not Available"),
        ])):
            _set_cell_border(cells[ci])
            cells[ci].width = width
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)

            if ci == 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                prio_lower = val.lower()
                if "immediate" in prio_lower:
                    r.font.color.rgb = RED
                    r.bold = True
                elif "short" in prio_lower or "30" in prio_lower:
                    r.font.color.rgb = AMBER
                    r.bold = True
                elif "long" in prio_lower or "90" in prio_lower:
                    r.font.color.rgb = BLUE
                elif "monitor" in prio_lower:
                    r.font.color.rgb = GREEN

    doc.add_paragraph()


def _build_summary_table(doc, inspection_data):
    _add_heading(doc, "Negative Side ↔ Positive Side Correlation", level=3, color=BLUE)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ci, hdr in enumerate(["Pt #", "Damage Observed (Negative Side)", "Pt #", "Source / Root Cause (Positive Side)"]):
        cell = tbl.rows[0].cells[ci]
        _set_cell_bg(cell, BLUE_DARK)
        _set_cell_border(cell)
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9)

    neg = inspection_data.summary_negative
    pos = inspection_data.summary_positive
    for n, p in zip(neg, pos):
        row = tbl.add_row()
        c = row.cells
        for ci, val in enumerate([n.get("point",""), n.get("description",""),
                                   p.get("point",""), p.get("description","")]):
            _set_cell_border(c[ci])
            c[ci].paragraphs[0].add_run(val).font.size = Pt(9)

    doc.add_paragraph()


def _build_conclusion(doc, analysis):
    doc.add_page_break()
    _add_heading(doc, "6. Conclusion", level=1)

    conclusion = analysis.get("conclusion", "Not Available")
    _add_paragraph(doc, conclusion, size=11, space_before=4, space_after=12)

    # Disclaimer
    _add_heading(doc, "Disclaimer", level=3, color=GRAY)
    disclaimer = (
        "This report has been prepared based on a visual inspection and thermal imaging survey conducted on the "
        "date stated herein. The findings are based on conditions observable at the time of inspection. "
        "This report does not constitute a structural engineering assessment. Recommendations should be "
        "implemented by qualified waterproofing contractors and verified by a licensed engineer where required."
    )
    _add_paragraph(doc, disclaimer, size=10, italic=True, color=GRAY)

    doc.add_paragraph()
    _add_paragraph(doc, "— End of Report —", bold=True, color=BLUE_DARK,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12)
